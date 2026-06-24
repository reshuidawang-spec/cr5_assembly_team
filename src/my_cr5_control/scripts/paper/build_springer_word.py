#!/usr/bin/env python3
"""Build a readable Springer-style Chinese Word manuscript from the Markdown draft."""

from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DOCS = ROOT / "paper_workspace/docs"
WORD = ROOT / "paper_workspace/word"
SOURCE = DOCS / "Q2_PAPER_CHINESE_DRAFT_V1.md"
PREPARED = WORD / "Q2_PAPER_CHINESE_SPRINGER_WORD_SOURCE.md"
RAW_DOCX = WORD / "Q2_PAPER_CHINESE_SPRINGER_RAW.docx"
FINAL_DOCX = WORD / "Q2_PAPER_CHINESE_SPRINGER_FORMATTED.docx"
PANDOC = ROOT / "paper_workspace/tools/pandoc-3.9.0.2/bin/pandoc"
BIB = ROOT / "paper_workspace/manuscript/latex/references.bib"
CSL = WORD / "styles/springer-basic-brackets.csl"

CITATIONS = {
    "（Ostrowska 等, 2022；Quigley 等, 2024；Jin 等, 2025）":
        "[@ostrowska2022correction; @quigley2024metrology; @jin2025review]",
    "（Sucan 等, 2012；Gammell 等, 2014, 2020；Strub 等, 2022）":
        "[@sucan2012ompl; @gammell2014informed; @gammell2020bitstar; @strub2022aitstar]",
    "（Sucan 等, 2012；Zhang 等, 2025）":
        "[@sucan2012ompl; @zhang2025samplingreview]",
    "（Malvido Fresnillo 等, 2023）": "[@malvido2023moveit]",
    "（Gammell 等, 2014, 2020；Strub 等, 2022；Yang 等, 2023）":
        "[@gammell2014informed; @gammell2020bitstar; @strub2022aitstar; @yang2023eplprm]",
    "（Yang 等, 2023）": "[@yang2023eplprm]",
    "（Wang 等, 2024a；Wang 等, 2024b）":
        "[@wang2024directionrrt; @wang2024heuristicmanipulator]",
    "（Fang 和 Ding, 2022；Yan 等, 2022；Li 等, 2024；Chen 等, 2024）":
        "[@fang2022activevisual; @yan2022twomoduleinspection; @li2024fiveaxis; @chen2024geodesicmeasurement]",
    "（Sliusarenko 等, 2023）": "[@sliusarenko2023constantprobe]",
    "（Ostrowska 等, 2022；Wu 等, 2023；Quigley 等, 2024）":
        "[@ostrowska2022correction; @wu2023joint; @quigley2024metrology]",
}

TABLE_CAPTIONS = (
    [(f"**表 {index}**", f"tab-{index}") for index in range(1, 9)]
    + [(f"**表 A{index}**", f"tab-A{index}") for index in range(1, 4)]
)

PROSE_REPLACEMENTS = {
    "**Target Format:** Springer Nature Journal Article Template (`sn-jnl`, Version 3.1, December 2024)\n\n": "",
    "六个测点定义如下，其中": "六个固定测点定义见[表 1](#tab-1)，其中",
    "该构造保留了方法的可解释性：": "固定目标构造见[表 2](#tab-2)。该构造保留了方法的可解释性：",
    "下表汇总复现实验所需的主要参数。": "[表 3](#tab-3)汇总复现实验所需的主要参数。",
    "预试验中，`FMT/BFMT`": "主对比实验配置汇总于[表 4](#tab-4)。预试验中，`FMT/BFMT`",
    "表 5 列出 `simple` 规则测试集上的结果。": "[表 5](#tab-5)列出 `simple` 规则测试集上的结果。",
    "表 6 给出 STL 工件 `v2/WS119` 上的结果。": "[表 6](#tab-6)给出 STL 工件 `v2/WS119` 上的结果。",
    "表 7 中，": "[表 7](#tab-7)中，",
    "下表报告成功率、耗时分位数、预算耗尽率和直接规划回退比例。":
        "[表 8](#tab-8)报告成功率、耗时分位数、预算耗尽率和直接规划回退比例。",
    "图 6 将导出的末端轨迹": "[图 6](#fig-6)将导出的末端轨迹",
    "列于附录表 A1": "列于[附录表 A1](#tab-A1)",
    "汇总于附录表 A2": "汇总于[附录表 A2](#tab-A2)",
    "见附录表 A3": "见[附录表 A3](#tab-A3)",
    "在表 A2 中，": "在[附录表 A2](#tab-A2)中，",
}


def prepare_source() -> None:
    """Prepare source."""
    text = SOURCE.read_text(encoding="utf-8")
    text = text.split("## References / 参考文献", 1)[0].rstrip() + "\n"
    for old, new in CITATIONS.items():
        text = text.replace(old, new)
    for old, new in PROSE_REPLACEMENTS.items():
        text = text.replace(old, new)
    for algorithm in ("RRT*", "BIT*", "AIT*"):
        text = text.replace(algorithm, algorithm[:-1] + r"\*")
    residual = re.findall(r"（[^）]*(?:19|20)\d{2}[^）]*）", text)
    if residual:
        raise RuntimeError(f"Unconverted author-year citation(s): {residual}")

    figure_index = 0
    table_index = 0
    out_lines: list[str] = []
    for line in text.splitlines():
        if line.startswith("**图 "):
            figure_index += 1
            out_lines.append(f"[]{{#fig-{figure_index}}}")
        if table_index < len(TABLE_CAPTIONS) and line.startswith(TABLE_CAPTIONS[table_index][0]):
            anchor = TABLE_CAPTIONS[table_index][1]
            table_index += 1
            out_lines.append(f"[]{{#{anchor}}}")
        out_lines.append(line)
    if figure_index != 7 or table_index != len(TABLE_CAPTIONS):
        raise RuntimeError(f"Expected 7 figures and 11 tables; found {figure_index} figures and {table_index} tables")
    out_lines.extend(["", "## 参考文献", "", "::: {#refs}", ":::", ""])
    PREPARED.write_text("\n".join(out_lines), encoding="utf-8")


def run_pandoc() -> None:
    """Run pandoc."""
    subprocess.run(
        [
            str(PANDOC),
            str(PREPARED),
            "--from=markdown+tex_math_single_backslash+tex_math_dollars+pipe_tables",
            "--to=docx",
            "--standalone",
            "--citeproc",
            f"--bibliography={BIB}",
            f"--csl={CSL}",
            "--metadata=link-citations:true",
            f"--resource-path={DOCS}:{ROOT / 'paper_workspace/manuscript/latex/figures'}",
            f"--output={RAW_DOCX}",
        ],
        check=True,
    )


def set_font(style, latin: str, east_asia: str, size: float, bold: bool = False) -> None:
    """Set font."""
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.bold = bold
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)


def shade_cell(cell, fill: str) -> None:
    """Shade cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 70, start: int = 90, bottom: int = 70, end: int = 90) -> None:
    """Set cell margins."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def append_page_field(paragraph) -> None:
    """Append page field."""
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    for element, value in (("w:fldChar", ("w:fldCharType", "begin")), ("w:instrText", None),
                           ("w:fldChar", ("w:fldCharType", "separate")), ("w:t", None),
                           ("w:fldChar", ("w:fldCharType", "end"))):
        node = OxmlElement(element)
        if value:
            node.set(qn(value[0]), value[1])
        elif element == "w:instrText":
            node.set(qn("xml:space"), "preserve")
            node.text = " PAGE "
        elif element == "w:t":
            node.text = "1"
        run._r.append(node)


def add_equation_number(paragraph, number: int) -> None:
    """Add equation number."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.tab_stops.clear_all()
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(8.0), WD_TAB_ALIGNMENT.CENTER)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(16.0), WD_TAB_ALIGNMENT.RIGHT)
    left = OxmlElement("w:r")
    left.append(OxmlElement("w:tab"))
    children = list(paragraph._p)
    insertion_index = 1 if children and children[0].tag == qn("w:pPr") else 0
    paragraph._p.insert(insertion_index, left)
    right = OxmlElement("w:r")
    right.append(OxmlElement("w:tab"))
    number_text = OxmlElement("w:t")
    number_text.text = f"({number})"
    right.append(number_text)
    paragraph._p.append(right)


def format_docx() -> tuple[int, int, int]:
    """Format docx."""
    doc = Document(RAW_DOCX)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    set_font(normal, "Times New Roman", "宋体", 10)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(5)
    set_font(doc.styles["Title"], "Times New Roman", "黑体", 16, True)
    doc.styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.styles["Title"].paragraph_format.space_after = Pt(12)
    set_font(doc.styles["Heading 1"], "Times New Roman", "黑体", 12, True)
    set_font(doc.styles["Heading 2"], "Times New Roman", "黑体", 11, True)
    for style_name in ("Heading 1", "Heading 2"):
        style = doc.styles[style_name]
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    caption = doc.styles.add_style("Formal Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption.base_style = normal
    set_font(caption, "Times New Roman", "宋体", 9)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_with_next = True
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(7)
    equation = doc.styles.add_style("Equation Formal", WD_STYLE_TYPE.PARAGRAPH)
    equation.base_style = normal
    set_font(equation, "Cambria Math", "Cambria Math", 10)
    equation.paragraph_format.space_before = Pt(4)
    equation.paragraph_format.space_after = Pt(4)
    references = doc.styles.add_style("References Formal", WD_STYLE_TYPE.PARAGRAPH)
    references.base_style = normal
    references.paragraph_format.first_line_indent = Cm(-0.6)
    references.paragraph_format.left_indent = Cm(0.6)
    references.paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.text = "Difficulty-Adaptive Informed Guide Sampling for Constrained Motion Planning"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(8)
        run.font.italic = True
    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    append_page_field(footer)
    for run in footer.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)

    equations = 0
    figures = 0
    in_references = False
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        original_style = paragraph.style.name
        if index == 0:
            paragraph.style = doc.styles["Title"]
        elif original_style == "Heading 2":
            paragraph.style = doc.styles["Heading 1"]
        elif original_style == "Heading 3":
            paragraph.style = doc.styles["Heading 2"]
        if text == "参考文献":
            in_references = True
        elif in_references and text:
            paragraph.style = references
        if re.match(r"^图 [1-7]\s+", text) or re.match(r"^表 (?:[1-8]|A[1-3])\s+", text):
            paragraph.style = caption
            if re.match(r"^图 [1-7]\s+", text):
                figures += 1
        if paragraph._p.xpath(".//w:drawing"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(8)
        if paragraph._p.xpath(".//m:oMathPara"):
            equations += 1
            paragraph.style = equation
            add_equation_number(paragraph, equations)

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        first_row = table.rows[0]
        first_row._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                if row_index == 0:
                    shade_cell(cell, "E7E6E6")
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.line_spacing = 1.05
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
                        run.font.size = Pt(8.5)
                        if row_index == 0:
                            run.font.bold = True
    doc.core_properties.title = "面向机器人接触测量受限路径规划的难度自适应引导采样方法"
    doc.core_properties.subject = "Springer-style Chinese manuscript draft"
    doc.save(FINAL_DOCX)
    return equations, figures, len(doc.tables)


def main() -> None:
    """Main."""
    WORD.mkdir(parents=True, exist_ok=True)
    prepare_source()
    run_pandoc()
    equations, figures, tables = format_docx()
    print(f"Generated: {FINAL_DOCX}")
    print(f"Display equations numbered: {equations}; figures: {figures}; tables: {tables}")


if __name__ == "__main__":
    main()
